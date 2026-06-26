from __future__ import annotations

import json
import re
from html import escape
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "assets"
DATA_DIR = ASSETS / "data"
CV_ASSETS = ASSETS / "CV"

JSON_FILES = {
    "profile": DATA_DIR / "profile.json",
    "publications": DATA_DIR / "publications.json",
    "conferences": DATA_DIR / "conferences.json",
    "awards": DATA_DIR / "awards.json",
    "patents": DATA_DIR / "patents.json",
    "cv": DATA_DIR / "cv.json",
    "links": DATA_DIR / "links.json",
    "projects": DATA_DIR / "projects.json",
}

# Keep the CV output aligned with the approved conference classification without
# modifying the Excel-generated JSON source in this cleanup phase.
ORAL_PRESENTATION_IDS = {"conf_001", "conf_002", "conf_003", "conf_005"}


def load_json(path: Path) -> Any:
    if not path.exists():
        raise FileNotFoundError(
            f"Missing {path.relative_to(ROOT)}. Run 'python tools\\generate_site_data.py' first."
        )
    return json.loads(path.read_text(encoding="utf-8"))


def load_site_data() -> dict[str, Any]:
    return {name: load_json(path) for name, path in JSON_FILES.items()}


def localized(value: Any, lang: str, fallback: str = "") -> str:
    if value is None:
        return fallback
    if isinstance(value, dict):
        return str(value.get(lang) or value.get("en") or value.get("zh") or fallback)
    return str(value)


def field(item: dict[str, Any], base: str, lang: str, fallback: str = "") -> str:
    return str(item.get(f"{base}_{lang}") or item.get(f"{base}_en") or item.get(base) or fallback or "")


def clean_text(text: Any) -> str:
    if text is None:
        return ""
    return re.sub(r"\s+", " ", str(text).replace("H. Sun", "Haolan Sun")).strip()


def sort_by_year(items: Iterable[dict[str, Any]]) -> list[dict[str, Any]]:
    def key(item: dict[str, Any]) -> tuple[int, str]:
        raw = item.get("year") or item.get("year_or_period") or item.get("date") or ""
        match = re.search(r"\d{4}", str(raw))
        year = int(match.group(0)) if match else -1
        return (-year, str(item.get("id") or item.get("item_id") or item.get("title_en") or ""))

    return sorted(items, key=key)


def is_oral_presentation(conf: dict[str, Any]) -> bool:
    return str(conf.get("id") or "") in ORAL_PRESENTATION_IDS


def publication_info(pub: dict[str, Any]) -> str:
    parts: list[str] = []
    journal = clean_text(pub.get("journal"))
    if journal:
        parts.append(journal)
    volume = clean_text(pub.get("volume"))
    issue = clean_text(pub.get("issue"))
    pages = clean_text(pub.get("pages"))
    year = clean_text(pub.get("year"))
    if volume and issue:
        parts.append(f"{volume}({issue})")
    elif volume:
        parts.append(volume)
    if pages:
        parts.append(pages)
    if year:
        parts.append(year)
    return ", ".join(parts)


def format_publication(pub: dict[str, Any], lang: str) -> str:
    title = clean_text(field(pub, "title", lang))
    authors = clean_text(pub.get("authors"))
    info = publication_info(pub)
    link = clean_text(pub.get("doi") or pub.get("link"))
    parts = [part for part in [title, authors, info, f"Link: {link}" if link else ""] if part]
    return ". ".join(parts)


def format_conference(conf: dict[str, Any], lang: str) -> str:
    title = clean_text(field(conf, "title", lang))
    authors = clean_text(conf.get("authors"))
    conference = clean_text(conf.get("conference"))
    location = clean_text(conf.get("location"))
    year = clean_text(conf.get("year"))
    venue = ", ".join(part for part in [conference, location, year] if part)
    link = clean_text(conf.get("link"))
    parts = [part for part in [title, authors, venue, f"Link: {link}" if link else ""] if part]
    return ". ".join(parts)


def format_award(award: dict[str, Any], lang: str) -> str:
    year = clean_text(award.get("year"))
    name = clean_text(field(award, "name", lang))
    level = clean_text(award.get("level"))
    org = clean_text(award.get("organization"))
    desc = clean_text(field(award, "description", lang))
    meta = " | ".join(part for part in [level, org, desc] if part)
    return " | ".join(part for part in [year, name, meta] if part)


def patent_status_label(patent: dict[str, Any], lang: str) -> str:
    status = clean_text(patent.get("status"))
    if lang == "zh":
        mapping = {
            "Granted invention patent": "授权发明专利",
            "Published invention patent application": "发明专利申请",
            "Utility model patent": "实用新型专利",
        }
        return mapping.get(status, status)
    return status


def format_patent(patent: dict[str, Any], lang: str) -> str:
    title = clean_text(field(patent, "title", lang))
    inventors = clean_text(patent.get("inventors"))
    applicant = clean_text(patent.get("applicant"))
    patent_no = clean_text(patent.get("patent_number"))
    app_no = clean_text(patent.get("application_number"))
    status = patent_status_label(patent, lang)
    app_date = clean_text(patent.get("application_date_display") or patent.get("application_date"))
    grant_date = clean_text(
        patent.get("grant_or_publication_date_display") or patent.get("grant_or_publication_date")
    )
    number = " / ".join(part for part in [patent_no, app_no] if part)
    if lang == "zh":
        pieces = [
            title,
            f"发明人: {inventors}" if inventors else "",
            f"申请人/专利权人: {applicant}" if applicant else "",
            f"专利/申请号: {number}" if number else "",
            f"状态: {status}" if status else "",
            f"申请日: {app_date}" if app_date else "",
            f"授权/公布日: {grant_date}" if grant_date else "",
        ]
    else:
        pieces = [
            title,
            f"Inventors: {inventors}" if inventors else "",
            f"Patentee / Applicant: {applicant}" if applicant else "",
            f"Patent No. / Application No.: {number}" if number else "",
            f"Status: {status}" if status else "",
            f"Application date: {app_date}" if app_date else "",
            f"Grant / publication date: {grant_date}" if grant_date else "",
        ]
    return ". ".join(part for part in pieces if part)


def patent_summary(patents: list[dict[str, Any]], lang: str) -> str:
    total = len(patents)
    invention = sum(1 for p in patents if p.get("patent_type") in {"granted_invention", "application_invention"})
    granted = sum(1 for p in patents if p.get("patent_type") == "granted_invention")
    utility = sum(1 for p in patents if p.get("patent_type") == "utility_model")
    if lang == "zh":
        return f"{total} 项专利 | {invention} 项发明专利 | {granted} 项授权发明 | {utility} 项实用新型"
    return f"{total} total patents | {invention} invention patents | {granted} granted inventions | {utility} utility model patents"


def format_cv_item(item: dict[str, Any], lang: str) -> str:
    period = clean_text(item.get("year_or_period"))
    title = clean_text(field(item, "title", lang))
    org = clean_text(item.get("organization"))
    desc = clean_text(field(item, "description", lang))
    return " | ".join(part for part in [period, title, org, desc] if part)


def cv_items(cv_data: dict[str, Any], section: str, lang: str, limit: int | None = None) -> list[str]:
    items = sorted(cv_data.get(section, []), key=lambda x: x.get("order") or 999)
    out = [format_cv_item(item, lang) for item in items]
    out = [item for item in out if item]
    return out[:limit] if limit else out


def build_sections(data: dict[str, Any], lang: str, academic: bool) -> list[dict[str, list[str] | str]]:
    cv_data = data["cv"]
    publications = sort_by_year(data["publications"])
    conferences = sort_by_year(data["conferences"])
    awards = sort_by_year(data["awards"])
    patents = sort_by_year(data["patents"])

    labels = {
        "en": {
            "education": "Education",
            "research": "Research Experience",
            "skills": "Technical Skills",
            "projects": "Projects" if academic else "Selected Projects",
            "publications": "Publications" if academic else "Selected Publications",
            "oral": "Conferences - Oral Presentations",
            "poster": "Conferences - Poster Presentations",
            "conferences": "Representative Conferences",
            "patents": "Patents",
            "awards": "Awards",
        },
        "zh": {
            "education": "教育经历",
            "research": "科研经历",
            "skills": "技术技能",
            "projects": "项目经历" if academic else "精选项目",
            "publications": "论文发表" if academic else "代表论文",
            "oral": "学术会议 - 口头报告",
            "poster": "学术会议 - 海报展示",
            "conferences": "代表会议",
            "patents": "专利",
            "awards": "奖励荣誉",
        },
    }[lang]

    sections: list[dict[str, list[str] | str]] = [
        {"title": labels["education"], "items": cv_items(cv_data, "education", lang)},
        {"title": labels["research"], "items": cv_items(cv_data, "research_experience", lang)},
        {"title": labels["skills"], "items": cv_items(cv_data, "technical_skills", lang)},
        {"title": labels["projects"], "items": cv_items(cv_data, "selected_projects", lang, None if academic else 3)},
    ]

    if academic:
        sections.append({"title": labels["publications"], "items": [format_publication(pub, lang) for pub in publications]})
        oral = [conf for conf in conferences if is_oral_presentation(conf)]
        poster = [conf for conf in conferences if not is_oral_presentation(conf)]
        sections.append({"title": labels["oral"], "items": [format_conference(conf, lang) for conf in oral]})
        sections.append({"title": labels["poster"], "items": [format_conference(conf, lang) for conf in poster]})
        sections.append(
            {
                "title": labels["patents"],
                "items": [patent_summary(patents, lang)] + [format_patent(patent, lang) for patent in patents],
            }
        )
        sections.append({"title": labels["awards"], "items": [format_award(award, lang) for award in awards]})
    else:
        selected_pubs = [pub for pub in publications if pub.get("selected")]
        if not selected_pubs:
            selected_pubs = publications[:5]
        sections.append(
            {"title": labels["publications"], "items": [format_publication(pub, lang) for pub in selected_pubs[:5]]}
        )
        selected_confs = cv_items(cv_data, "selected_conferences", lang, 5)
        sections.append({"title": labels["conferences"], "items": selected_confs})
        representative_patents = cv_items(cv_data, "patent_overview", lang, 5)
        sections.append({"title": labels["patents"], "items": representative_patents})
        sections.append({"title": labels["awards"], "items": cv_items(cv_data, "awards", lang)})

    return sections


def build_cv_payloads(site_data: dict[str, Any]) -> list[dict[str, Any]]:
    profile = site_data["profile"]
    cv_data = site_data["cv"]
    profile_entry = cv_data.get("profile", [{}])[0]
    name_en = localized(profile.get("name"), "en", "Longlong Li")
    name_zh = localized(profile.get("name"), "zh", "李龙龙")
    position_en = localized(profile.get("position"), "en", "PhD Student")
    position_zh = localized(profile.get("position"), "zh", "博士研究生")
    affiliation_en = localized(profile.get("affiliation"), "en")
    affiliation_zh = localized(profile.get("affiliation"), "zh")
    location_en = localized(profile.get("location"), "en")
    location_zh = localized(profile.get("location"), "zh")
    email = localized(profile.get("email"), "en", "lilonglong2000@jnu.ac.kr")
    homepage = localized(profile.get("homepage"), "en", "longlongli.com")

    common = {
        "email": email,
        "homepage": homepage,
    }

    definitions = [
        {
            "filename_base": "Longlong_Li_Academic_CV_EN",
            "language": "en",
            "academic": True,
            "name": name_en,
            "title": "Academic CV",
            "position": f"{position_en} | BioMEMS / Biosensors / Engineered Cardiac Tissue Monitoring",
            "affiliation": affiliation_en,
            "location": location_en,
            "header_note": "Comprehensive academic version for research collaboration, postdoctoral applications, and academic review.",
            "summary": [clean_text(field(profile_entry, "description", "en"))],
        },
        {
            "filename_base": "Longlong_Li_Academic_CV_ZH",
            "language": "zh",
            "academic": True,
            "name": name_zh,
            "title": "学术简历",
            "position": f"{position_zh} | BioMEMS / 生物传感器 / 工程化心脏组织监测",
            "affiliation": affiliation_zh,
            "location": location_zh,
            "header_note": "面向学术合作、博士后申请和成果归档的完整版本。",
            "summary": [clean_text(field(profile_entry, "description", "zh"))],
        },
        {
            "filename_base": "Longlong_Li_Resume_EN",
            "language": "en",
            "academic": False,
            "name": name_en,
            "title": "2-page Resume",
            "position": f"{position_en} | BioMEMS / Biosensors / Engineered Cardiac Tissue Monitoring",
            "affiliation": affiliation_en,
            "location": location_en,
            "header_note": "Condensed version for HR screening, R&D roles, and industry-facing applications.",
            "summary": [clean_text(field(profile_entry, "description", "en"))],
        },
        {
            "filename_base": "Longlong_Li_Resume_ZH",
            "language": "zh",
            "academic": False,
            "name": name_zh,
            "title": "2 页版求职简历",
            "position": f"{position_zh} | BioMEMS / 生物传感器 / 工程化心脏组织监测",
            "affiliation": affiliation_zh,
            "location": location_zh,
            "header_note": "面向 HR 初筛、研发岗位和产业应用场景的精简版本。",
            "summary": [clean_text(field(profile_entry, "description", "zh"))],
        },
    ]

    payloads: list[dict[str, Any]] = []
    for item in definitions:
        payload = {**common, **item}
        payload["sections"] = build_sections(site_data, payload["language"], payload["academic"])
        payloads.append(payload)
    return payloads


def set_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)


def add_bottom_rule(paragraph) -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), "B6C4D5")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def style_document(doc: Document, east_asia_font: str, compact: bool) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5 if compact else 10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.08 if compact else 1.12

    title = doc.styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(22 if compact else 24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(18, 54, 95)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)

    heading = doc.styles["Heading 1"]
    heading.font.name = "Arial"
    heading.font.size = Pt(11.5)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor(29, 78, 137)
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    heading.paragraph_format.space_before = Pt(8 if compact else 10)
    heading.paragraph_format.space_after = Pt(3)


def add_header(doc: Document, data: dict[str, Any], east_asia_font: str, compact: bool) -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(data["name"])
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(1)
    for text in (data["position"], data["affiliation"]):
        r = p2.add_run(text + "  ")
        r.bold = True
        r.font.size = Pt(10 if compact else 10.5)
        r.font.name = "Arial"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(4)
    meta = [data["location"], data["email"], f"Homepage: {data['homepage']}"]
    for idx, text in enumerate(meta):
        r = p3.add_run(text)
        r.font.size = Pt(8.8 if compact else 9.2)
        r.font.name = "Arial"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
        if idx < len(meta) - 1:
            sep = p3.add_run(" | ")
            sep.font.size = Pt(8.8 if compact else 9.2)
            sep.font.name = "Arial"
            sep._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(8)
    note.paragraph_format.left_indent = Inches(0.03)
    note.paragraph_format.right_indent = Inches(0.03)
    add_bottom_rule(note)
    r = note.add_run(data["header_note"])
    r.italic = True
    r.font.size = Pt(9 if compact else 9.4)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)


def add_bullet_group(doc: Document, title: str, items: list[str], east_asia_font: str, compact: bool) -> None:
    if not items:
        return
    h = doc.add_paragraph(style="Heading 1")
    r = h.add_run(title)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)

    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        p.paragraph_format.space_after = Pt(1 if compact else 2)
        bullet = p.add_run("- ")
        bullet.bold = True
        bullet.font.name = "Arial"
        bullet._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
        body = p.add_run(item)
        body.font.name = "Arial"
        body._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)


def add_footer(doc: Document, label: str) -> None:
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(label)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(86, 99, 117)


def build_docx(data: dict[str, Any]) -> Path:
    east_asia_font = "SimSun" if data["language"] == "zh" else "Arial"
    compact = not data["academic"]
    doc = Document()
    set_page(doc)
    style_document(doc, east_asia_font, compact)
    add_header(doc, data, east_asia_font, compact)
    profile_title = "Profile" if data["language"] == "en" else "个人简介"
    add_bullet_group(doc, profile_title, data["summary"], east_asia_font, compact)
    for section in data["sections"]:
        add_bullet_group(doc, str(section["title"]), list(section["items"]), east_asia_font, compact)
    add_footer(doc, data["filename_base"])
    path = CV_ASSETS / f"{data['filename_base']}.docx"
    doc.save(path)
    return path


def pdf_paragraph(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(escape(text), style)


def build_pdf(data: dict[str, Any]) -> Path:
    out = CV_ASSETS / f"{data['filename_base']}.pdf"
    compact = not data["academic"]
    font_name = "STSong-Light" if data["language"] == "zh" else "Helvetica"
    title_font = "STSong-Light" if data["language"] == "zh" else "Helvetica-Bold"

    doc = SimpleDocTemplate(
        str(out),
        pagesize=letter,
        leftMargin=0.72 * inch,
        rightMargin=0.72 * inch,
        topMargin=0.64 * inch,
        bottomMargin=0.58 * inch,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleStyle",
        parent=styles["Title"],
        fontName=title_font,
        fontSize=18 if compact else 20,
        leading=21 if compact else 24,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#12365F"),
        spaceAfter=4,
    )
    meta_style = ParagraphStyle(
        "MetaStyle",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=8.4 if compact else 8.8,
        leading=10.5 if compact else 11.5,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#334155"),
        spaceAfter=2,
    )
    note_style = ParagraphStyle(
        "NoteStyle",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=8.6 if compact else 9,
        leading=11 if compact else 12,
        textColor=colors.HexColor("#243143"),
        backColor=colors.HexColor("#F4F7FB"),
        borderColor=colors.HexColor("#B6C4D5"),
        borderWidth=0.4,
        borderPadding=5,
        spaceAfter=7,
    )
    heading_style = ParagraphStyle(
        "HeadingStyle",
        parent=styles["Heading1"],
        fontName=title_font,
        fontSize=10.4 if compact else 11.2,
        leading=12.5 if compact else 13.8,
        textColor=colors.HexColor("#1D4E89"),
        spaceBefore=5 if compact else 7,
        spaceAfter=2,
    )
    bullet_style = ParagraphStyle(
        "BulletStyle",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=8.5 if compact else 9,
        leading=10.5 if compact else 11.5,
        leftIndent=12,
        firstLineIndent=-9,
        spaceAfter=1,
        textColor=colors.black,
    )

    story = [
        pdf_paragraph(data["name"], title_style),
        pdf_paragraph(data["position"], meta_style),
        pdf_paragraph(data["affiliation"], meta_style),
        pdf_paragraph(f"{data['location']} | {data['email']} | Homepage: {data['homepage']}", meta_style),
        Spacer(1, 2),
        pdf_paragraph(data["header_note"], note_style),
    ]

    profile_title = "Profile" if data["language"] == "en" else "个人简介"
    story.append(pdf_paragraph(profile_title, heading_style))
    for item in data["summary"]:
        story.append(pdf_paragraph(f"- {item}", bullet_style))

    for section in data["sections"]:
        story.append(pdf_paragraph(str(section["title"]), heading_style))
        for item in section["items"]:
            story.append(pdf_paragraph(f"- {item}", bullet_style))

    doc.build(story)
    return out


def print_counts(site_data: dict[str, Any]) -> None:
    conferences = site_data["conferences"]
    source_oral = sum(1 for item in conferences if clean_text(item.get("presentation_type")).lower() == "oral")
    oral = sum(1 for item in conferences if is_oral_presentation(item))
    poster = len(conferences) - oral
    patents = site_data["patents"]
    print("Loaded website JSON data:")
    print(f"- publications: {len(site_data['publications'])}")
    print(f"- conferences: {len(conferences)} ({oral} oral, {poster} poster in generated CV)")
    if source_oral != oral:
        print(f"  note: source JSON marks {source_oral} records as oral; CV output uses the approved 4-record oral classification.")
    print(f"- awards: {len(site_data['awards'])}")
    print(f"- patents: {len(patents)}")


def main() -> None:
    CV_ASSETS.mkdir(parents=True, exist_ok=True)
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    site_data = load_site_data()
    print_counts(site_data)
    payloads = build_cv_payloads(site_data)
    print("Generated CV documents:")
    for data in payloads:
        docx_path = build_docx(data)
        pdf_path = build_pdf(data)
        print(f"- {docx_path.relative_to(ROOT)}")
        print(f"- {pdf_path.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
